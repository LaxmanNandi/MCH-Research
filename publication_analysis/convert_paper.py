"""
Convert MCH paper to DOCX and PDF, then create ZIP
"""

import os
import sys
import zipfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown2
import re

# Paths
BASE_DIR = "C:/Users/barla/mch_experiments/publication_analysis"
MD_FILE = os.path.join(BASE_DIR, "MCH_arXiv_Paper.md")
DOCX_FILE = os.path.join(BASE_DIR, "MCH_arXiv_Paper.docx")
PDF_FILE = os.path.join(BASE_DIR, "MCH_arXiv_Paper.pdf")
ZIP_FILE = os.path.join(BASE_DIR, "MCH_Publication_Package.zip")

def read_markdown():
    """Read the markdown file."""
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        return f.read()

def restart_numbering(paragraph):
    """Force restart numbering for a list paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()

    # Add ilvl (indentation level)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.insert(0, ilvl)

    # Add numId
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(numId)

def create_docx(md_content):
    """Create DOCX from markdown content with proper list numbering."""
    print("Creating DOCX...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0
    in_numbered_list = False
    last_was_numbered = False
    list_counter = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            last_was_numbered = False
            i += 1
            continue

        # Check if this is a numbered list item
        is_numbered = bool(re.match(r'^\d+\. ', line))

        # Detect start of new numbered list
        start_new_list = is_numbered and not last_was_numbered

        # Main title (# )
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:]
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_was_numbered = False

        # Section headers (## )
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            last_was_numbered = False

        # Subsection headers (### )
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            last_was_numbered = False

        # Horizontal rule
        elif line == '---':
            doc.add_paragraph('_' * 50)
            last_was_numbered = False

        # Code blocks
        elif line.startswith('```'):
            # Skip the opening ```
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # Add code as formatted paragraph
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            last_was_numbered = False

        # Table detection
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            # Parse table
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            # Skip header separator
            if len(table_lines) > 1:
                headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]

                # Create table
                rows_data = []
                for tl in table_lines[2:]:  # Skip header and separator
                    if '|' in tl:
                        cells = [cell.strip() for cell in tl.split('|') if cell.strip()]
                        if cells:
                            rows_data.append(cells)

                if headers and rows_data:
                    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
                    table.style = 'Table Grid'

                    # Headers
                    for ci, header in enumerate(headers):
                        if ci < len(table.rows[0].cells):
                            table.rows[0].cells[ci].text = header

                    # Data rows
                    for ri, row_data in enumerate(rows_data):
                        for ci, cell_data in enumerate(row_data):
                            if ci < len(table.rows[ri + 1].cells):
                                table.rows[ri + 1].cells[ci].text = cell_data

                i = j - 1
            last_was_numbered = False

        # Bullet points
        elif line.startswith('- '):
            text = line[2:]
            # Handle bold within bullets
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            last_was_numbered = False

        # Numbered lists - preserve original numbers from markdown
        elif is_numbered:
            # Extract the number and text
            match = re.match(r'^(\d+)\. (.+)$', line)
            if match:
                num = match.group(1)
                text = match.group(2)
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

                # Add as regular paragraph with manual number
                p = doc.add_paragraph(f"{num}. {text}")

            last_was_numbered = True

        # Bold paragraph starts
        elif line.startswith('**') and ':**' in line:
            # Definition style
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
            p = doc.add_paragraph(text)
            last_was_numbered = False

        # Regular paragraphs
        else:
            # Remove markdown formatting
            text = line
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links

            if text:
                doc.add_paragraph(text)
            last_was_numbered = False

        i += 1

    doc.save(DOCX_FILE)
    print(f"  Saved: {DOCX_FILE}")

def create_pdf_from_html(md_content):
    """Create PDF using weasyprint."""
    print("Creating PDF...")

    try:
        from weasyprint import HTML, CSS

        # Convert markdown to HTML
        html_content = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])

        # Wrap in HTML document with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: 'Times New Roman', Times, serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    max-width: 8.5in;
                    margin: 1in;
                    padding: 0;
                }}
                h1 {{
                    font-size: 18pt;
                    text-align: center;
                    margin-bottom: 0.5in;
                }}
                h2 {{
                    font-size: 14pt;
                    margin-top: 0.3in;
                    border-bottom: 1px solid #333;
                }}
                h3 {{
                    font-size: 12pt;
                    margin-top: 0.2in;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 0.2in 0;
                    font-size: 10pt;
                }}
                th, td {{
                    border: 1px solid #333;
                    padding: 6px;
                    text-align: left;
                }}
                th {{
                    background-color: #f0f0f0;
                }}
                p {{
                    margin: 0.1in 0;
                    text-align: justify;
                }}
                ul, ol {{
                    margin-left: 0.3in;
                }}
                hr {{
                    border: none;
                    border-top: 1px solid #999;
                    margin: 0.2in 0;
                }}
                code {{
                    font-family: 'Courier New', monospace;
                    font-size: 10pt;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        HTML(string=full_html).write_pdf(PDF_FILE)
        print(f"  Saved: {PDF_FILE}")
        return True

    except Exception as e:
        print(f"  PDF creation failed: {e}")
        print("  Trying alternative method...")
        return False

def create_zip():
    """Create ZIP file of all publication files."""
    print("Creating ZIP package...")

    files_to_zip = [
        "MCH_arXiv_Paper.md",
        "MCH_arXiv_Paper.docx",
        "mch_complete_analysis_report.txt",
        "figure1_response_coherence.png",
        "figure2_effect_sizes.png",
        "figure3_vendor_tier.png",
        "mch_complete_dataset.json",
        "mch_data_README.md",
        "mch_validation_report.txt",
        "app.py",
        "requirements.txt",
        "STREAMLIT_README.md",
    ]

    # Add PDF if it exists
    if os.path.exists(PDF_FILE):
        files_to_zip.append("MCH_arXiv_Paper.pdf")

    with zipfile.ZipFile(ZIP_FILE, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for filename in files_to_zip:
            filepath = os.path.join(BASE_DIR, filename)
            if os.path.exists(filepath):
                zipf.write(filepath, filename)
                print(f"  Added: {filename}")
            else:
                print(f"  Skipped (not found): {filename}")

    print(f"  Saved: {ZIP_FILE}")

def main():
    print("=" * 60)
    print("CONVERTING PAPER TO DOCX, PDF, AND CREATING ZIP")
    print("=" * 60)

    # Read markdown
    md_content = read_markdown()
    print(f"Read {len(md_content)} characters from markdown")

    # Create DOCX
    create_docx(md_content)

    # Create PDF
    create_pdf_from_html(md_content)

    # Create ZIP
    create_zip()

    print("\n" + "=" * 60)
    print("CONVERSION COMPLETE!")
    print("=" * 60)
    print(f"\nFiles in {BASE_DIR}:")
    print("  - MCH_arXiv_Paper.md (original)")
    print("  - MCH_arXiv_Paper.docx")
    if os.path.exists(PDF_FILE):
        print("  - MCH_arXiv_Paper.pdf")
    print("  - MCH_Publication_Package.zip")

if __name__ == "__main__":
    main()
